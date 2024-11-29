import os
import uuid
import json
from datetime import datetime
from tkinter import Tk, Label, Button, Text, filedialog, StringVar, OptionMenu, END, messagebox, Frame
from tkinter.scrolledtext import ScrolledText
from dotenv import load_dotenv
from openai import OpenAI  # Beibehalten gemäß Benutzeranforderung
from docx import Document
import PyPDF2
from PIL import Image
import base64
import shutil
import logging

# =========================
# Konfiguration und Setup
# =========================

# Logging konfigurieren
logging.basicConfig(level=logging.DEBUG, filename='app.log',
                    format='%(asctime)s - %(levelname)s - %(message)s')

# Basisverzeichnis festlegen
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Umgebungsvariablen laden
load_dotenv(os.path.join(BASE_DIR, '.env'))
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')

if not OPENAI_API_KEY:
    logging.error("OpenAI API-Schlüssel nicht gefunden.")
    raise ValueError("OpenAI API key not found. Please set it in the .env file.")

# OpenAI-Client initialisieren
client = OpenAI(api_key=OPENAI_API_KEY)

# Verzeichnisse konfigurieren
CHAT_HISTORY_FOLDER = os.path.join(BASE_DIR, 'chat_history')
EXPORTED_CHATS_FOLDER = os.path.join(BASE_DIR, 'exported_chats')
PROMPTS_FOLDER = os.path.join(BASE_DIR, 'prompts')
ATTACHMENTS_FOLDER = os.path.join(BASE_DIR, 'attachments')

# Sicherstellen, dass alle notwendigen Verzeichnisse existieren
for folder in [CHAT_HISTORY_FOLDER, EXPORTED_CHATS_FOLDER, PROMPTS_FOLDER, ATTACHMENTS_FOLDER]:
    try:
        os.makedirs(folder, exist_ok=True)
        if os.path.exists(folder):
            logging.debug(f"Ordner '{folder}' existiert oder wurde erfolgreich erstellt.")
        else:
            messagebox.showerror("Fehler", f"Ordner '{folder}' konnte nicht erstellt werden.")
            logging.error(f"Ordner '{folder}' konnte nicht erstellt werden.")
    except Exception as e:
        messagebox.showerror("Fehler", f"Fehler beim Erstellen des Ordners '{folder}': {str(e)}")
        logging.error(f"Fehler beim Erstellen des Ordners '{folder}': {str(e)}")

# =========================
# Funktionen
# =========================

def get_prompts():
    """
    Lädt alle Prompts aus dem PROMPTS_FOLDER und gibt sie als Liste von Dictionaries zurück.
    """
    prompts = []
    try:
        prompt_files = [f for f in os.listdir(PROMPTS_FOLDER) if f.endswith('.txt')]
        logging.debug(f"Gefundene Prompt-Dateien: {prompt_files}")
    except FileNotFoundError:
        messagebox.showerror("Fehler", f"Der Ordner '{PROMPTS_FOLDER}' wurde nicht gefunden.")
        logging.error(f"Der Ordner '{PROMPTS_FOLDER}' wurde nicht gefunden.")
        return prompts

    for filename in prompt_files:
        prompt_path = os.path.join(PROMPTS_FOLDER, filename)
        try:
            with open(prompt_path, 'r', encoding='utf-8') as f:
                content = f.read()
                prompts.append({'name': filename, 'content': content})
                logging.debug(f"Prompt '{filename}' erfolgreich geladen.")
        except Exception as e:
            messagebox.showwarning("Warnung", f"Fehler beim Laden der Datei '{filename}': {str(e)}")
            logging.warning(f"Fehler beim Laden der Datei '{filename}': {str(e)}")
    return prompts

# =========================
# ChatbotApp Klasse
# =========================

class ChatbotApp:
    def __init__(self, master):
        self.master = master
        master.title("Lokaler Chatbot mit Tkinter")
        master.geometry("800x600")
        logging.debug("ChatbotApp wurde initialisiert.")

        # Model Selection
        self.model_frame = Frame(master)
        self.model_frame.pack(pady=5)
        
        self.model_label = Label(self.model_frame, text="Model:")
        self.model_label.pack(side='left')
        
        self.selected_model = StringVar(master)
        self.selected_model.set("gpt-4o-mini")  # Standardmodell setzen
        self.model_menu = OptionMenu(self.model_frame, self.selected_model, "gpt-4o-mini", "gpt-4o")
        self.model_menu.pack(side='left')

        # Attachment Frame
        self.attachment_frame = Frame(master)
        self.attachment_frame.pack(pady=5)
        
        self.attachment_button = Button(self.attachment_frame, text="Datei anhängen", command=self.attach_file)
        self.attachment_button.pack(side='left', padx=5)
        
        self.attachment_label = Label(self.attachment_frame, text="Keine Datei ausgewählt")
        self.attachment_label.pack(side='left')

        # Prompt Selection
        self.prompt_label = Label(master, text="Wähle ein Prompt:")
        self.prompt_label.pack()

        self.prompts = get_prompts()
        if not self.prompts:
            messagebox.showerror("Fehler", "Keine Prompts gefunden im 'prompts'-Ordner.")
            master.quit()

        self.selected_prompt = StringVar(master)
        self.selected_prompt.set(self.prompts[0]['name'])

        self.prompt_menu = OptionMenu(master, self.selected_prompt, *[p['name'] for p in self.prompts])
        self.prompt_menu.pack()

        # User Input
        self.user_input_label = Label(master, text="Deine Eingabe:")
        self.user_input_label.pack()

        self.user_input_text = Text(master, height=4, width=80)
        self.user_input_text.pack()

        # Send Button
        self.send_button = Button(master, text="Senden", command=self.start_chat)
        self.send_button.pack(pady=5)

        # Chat History Display
        self.chat_display = ScrolledText(master, state='disabled', height=20, width=80)
        self.chat_display.pack(pady=5)

        # Export Button
        self.export_button = Button(master, text="Chat exportieren", command=self.export_chat)
        self.export_button.pack(pady=5)

        # Initialize attachment storage
        self.current_attachment = None
        self.current_chat_id = None

    def attach_file(self):
        """
        Öffnet einen Datei-Dialog zur Auswahl einer unterstützten Datei und speichert sie im 'attachments'-Ordner.
        """
        file_path = filedialog.askopenfilename(
            filetypes=[
                ("Supported files", "*.pdf;*.png;*.jpg;*.jpeg"),
                ("PDF files", "*.pdf"),
                ("Image files", "*.png;*.jpg;*.jpeg")
            ]
        )
        
        if file_path:
            file_name = os.path.basename(file_path)
            destination_path = os.path.join(ATTACHMENTS_FOLDER, file_name)
            try:
                shutil.copy(file_path, destination_path)
                self.attachment_label.config(text=file_name)
                self.current_attachment = self.process_attachment(destination_path)
                logging.debug(f"Datei '{file_name}' wurde in '{ATTACHMENTS_FOLDER}' kopiert.")
            except Exception as e:
                messagebox.showerror("Fehler", f"Fehler beim Kopieren der Datei: {str(e)}")
                logging.error(f"Fehler beim Kopieren der Datei '{file_name}': {str(e)}")

    def process_attachment(self, file_path):
        """
        Verarbeitet die angehängte Datei je nach Dateityp.
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self.process_pdf(file_path)
        elif file_ext in ['.png', '.jpg', '.jpeg']:
            return self.process_image(file_path)
        else:
            messagebox.showerror("Fehler", "Unsupported file type")
            logging.error(f"Unsupported file type: {file_ext}")
            return None

    def process_pdf(self, file_path):
        """
        Extrahiert den Text aus einer PDF-Datei.
        """
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text
                    else:
                        logging.warning(f"Keine Textinhalte auf Seite {page_num + 1} der PDF '{file_path}'.")
                logging.debug(f"PDF '{file_path}' erfolgreich verarbeitet.")
                return {"type": "pdf", "content": text}
        except Exception as e:
            messagebox.showerror("Fehler", f"Failed to process PDF: {str(e)}")
            logging.error(f"Failed to process PDF '{file_path}': {str(e)}")
            return None

    def process_image(self, file_path):
        """
        Kodiert ein Bild in Base64.
        """
        try:
            with open(file_path, "rb") as image_file:
                encoded_string = base64.b64encode(image_file.read()).decode('utf-8')
                logging.debug(f"Bild '{file_path}' erfolgreich verarbeitet.")
                return {"type": "image", "content": encoded_string}
        except Exception as e:
            messagebox.showerror("Fehler", f"Failed to process image: {str(e)}")
            logging.error(f"Failed to process image '{file_path}': {str(e)}")
            return None

    def start_chat(self):
        """
        Startet den Chat, sendet die Eingabe an OpenAI und zeigt die Antwort an.
        """
        prompt_name = self.selected_prompt.get()
        user_input = self.user_input_text.get("1.0", END).strip()
        model = self.selected_model.get()

        if not user_input:
            messagebox.showwarning("Eingabe erforderlich", "Bitte gib eine Nachricht ein.")
            logging.warning("Benutzer hat versucht zu senden, ohne eine Eingabe zu tätigen.")
            return

        # Load selected prompt
        selected_prompt = None
        for prompt in self.prompts:
            if prompt['name'] == prompt_name:
                selected_prompt = prompt['content']
                break

        if not selected_prompt:
            messagebox.showerror("Fehler", "Ausgewählter Prompt nicht gefunden.")
            logging.error(f"Ausgewählter Prompt '{prompt_name}' nicht gefunden.")
            return

        # Generate unique chat ID if not exists
        if not self.current_chat_id:
            self.current_chat_id = str(uuid.uuid4())
            chat_folder = os.path.join(CHAT_HISTORY_FOLDER, self.current_chat_id)
            try:
                os.makedirs(chat_folder, exist_ok=True)
                logging.debug(f"Chat-Verzeichnis '{chat_folder}' erstellt.")
            except Exception as e:
                messagebox.showerror("Fehler", f"Fehler beim Erstellen des Chat-Verzeichnisses: {str(e)}")
                logging.error(f"Fehler beim Erstellen des Chat-Verzeichnisses '{chat_folder}': {str(e)}")
                return

        # Initialize chat messages
        messages = [
            {"role": "system", "content": selected_prompt},
            {"role": "user", "content": user_input}
        ]

        # Add attachment if present
        if self.current_attachment:
            if self.current_attachment["type"] == "pdf":
                messages.append({
                    "role": "user",
                    "content": f"Here's the content of the attached PDF:\n{self.current_attachment['content']}"
                })
            elif self.current_attachment["type"] == "image":
                messages.append({
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Here's an attached image:"},
                        {"type": "image_url", "image_url": f"data:image/jpeg;base64,{self.current_attachment['content']}"}
                    ]
                })

        # Initialize chat history
        chat_history = []
        timestamp = datetime.utcnow().isoformat()
        chat_history.append({'role': 'system', 'content': selected_prompt, 'timestamp': timestamp})
        chat_history.append({'role': 'user', 'content': user_input, 'timestamp': timestamp})

        # Send request to OpenAI API
        try:
            response = client.chat.completions.create(
                model=model,
                messages=messages
            )
            reply = response.choices[0].message.content
            chat_history.append({'role': 'assistant', 'content': reply, 'timestamp': datetime.utcnow().isoformat()})
            logging.debug("Antwort von OpenAI erfolgreich erhalten.")
        except Exception as e:
            reply = f"Ein Fehler ist aufgetreten: {str(e)}"
            chat_history.append({'role': 'assistant', 'content': reply, 'timestamp': datetime.utcnow().isoformat()})
            logging.error(f"Fehler bei der OpenAI-Anfrage: {str(e)}")

        # Save chat history
        history_path = os.path.join(CHAT_HISTORY_FOLDER, self.current_chat_id, 'history.json')
        try:
            with open(history_path, 'w', encoding='utf-8') as f:
                json.dump(chat_history, f, ensure_ascii=False, indent=4)
            logging.debug(f"Chat-Historie erfolgreich in '{history_path}' gespeichert.")
        except Exception as e:
            messagebox.showerror("Fehler", f"Fehler beim Speichern der Chat-Historie: {str(e)}")
            logging.error(f"Fehler beim Speichern der Chat-Historie '{history_path}': {str(e)}")
            return

        # Update chat display
        self.chat_display.configure(state='normal')
        self.chat_display.insert(END, f"Du: {user_input}\n")
        self.chat_display.insert(END, f"ChatGPT: {reply}\n\n")
        self.chat_display.configure(state='disabled')
        self.chat_display.see(END)

        # Clear input and attachment
        self.user_input_text.delete("1.0", END)
        self.current_attachment = None
        self.attachment_label.config(text="Keine Datei ausgewählt")
        logging.debug("Eingabefeld und Anhänge wurden zurückgesetzt.")

    def export_chat(self):
        """
        Exportiert die aktuelle Chat-Historie als DOCX-Datei.
        """
        if not self.current_chat_id:
            messagebox.showinfo("Keine Chats", "Es gibt keine Chat-Historie zum Exportieren.")
            logging.info("Export versucht ohne vorhandene Chat-Historie.")
            return

        history_file = os.path.join(CHAT_HISTORY_FOLDER, self.current_chat_id, 'history.json')
        if not os.path.exists(history_file):
            messagebox.showerror("Fehler", "Chat-Historie nicht gefunden.")
            logging.error(f"Chat-Historie-Datei '{history_file}' nicht gefunden.")
            return

        try:
            with open(history_file, 'r', encoding='utf-8') as f:
                chat_history = json.load(f)
            logging.debug(f"Chat-Historie aus '{history_file}' geladen.")
        except Exception as e:
            messagebox.showerror("Fehler", f"Fehler beim Laden der Chat-Historie: {str(e)}")
            logging.error(f"Fehler beim Laden der Chat-Historie '{history_file}': {str(e)}")
            return

        # Create DOCX document
        doc = Document()
        doc.add_heading(f'Chat Session: {self.current_chat_id}', 0)

        for entry in chat_history:
            if entry['role'] == 'system':
                continue
            elif entry['role'] == 'user':
                doc.add_paragraph(f"User: {entry['content']}", style='Intense Quote')
            elif entry['role'] == 'assistant':
                doc.add_paragraph(f"ChatGPT: {entry['content']}")

        exported_filename = f"{self.current_chat_id}.docx"
        exported_path = os.path.join(EXPORTED_CHATS_FOLDER, exported_filename)
        try:
            doc.save(exported_path)
            logging.debug(f"Chat wurde als '{exported_path}' gespeichert.")
        except Exception as e:
            messagebox.showerror("Fehler", f"Fehler beim Speichern des DOCX-Dokuments: {str(e)}")
            logging.error(f"Fehler beim Speichern des DOCX-Dokuments '{exported_path}': {str(e)}")
            return

        # Open file dialog for export
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=exported_filename,
            filetypes=[("Word Documents", "*.docx")]
        )
        
        if save_path:
            try:
                shutil.move(exported_path, save_path)
                messagebox.showinfo("Erfolgreich", f"Chat wurde nach {save_path} exportiert.")
                logging.debug(f"Chat '{self.current_chat_id}' erfolgreich nach '{save_path}' exportiert.")
            except Exception as e:
                messagebox.showerror("Fehler", f"Fehler beim Exportieren: {str(e)}")
                logging.error(f"Fehler beim Exportieren des Chats nach '{save_path}': {str(e)}")
        else:
            try:
                os.remove(exported_path)
                logging.debug(f"Temporäre Datei '{exported_path}' wurde gelöscht.")
            except Exception as e:
                messagebox.showwarning("Warnung", f"Fehler beim Löschen der temporären Datei: {str(e)}")
                logging.warning(f"Fehler beim Löschen der temporären Datei '{exported_path}': {str(e)}")

# =========================
# Hauptfunktion
# =========================

def main():
    root = Tk()
    app = ChatbotApp(root)
    root.mainloop()

if __name__ == '__main__':
    main()
